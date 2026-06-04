import openpyxl, json, sys
from docx import Document

output = {}

# 1. 解析xlsx
print("=== 解析xlsx ===")
wb = openpyxl.load_workbook(r'd:\ai_guide\景点景区旅游数据行为分析数据.xlsx', data_only=True)
sheets_data = {}
for name in wb.sheetnames:
    ws = wb[name]
    rows_list = []
    for row in ws.iter_rows(values_only=True):
        rows_list.append([str(c) if c is not None else '' for c in row])
    sheets_data[name] = rows_list
    print(f"Sheet [{name}]: {len(rows_list)} rows")
output['xlsx'] = sheets_data

# 2. 解析结构化数据集docx
print("\n=== 解析结构化数据集docx ===")
doc1 = Document(r'd:\ai_guide\灵山胜境 景点结构化数据集.docx')
struct_text = []
for para in doc1.paragraphs:
    t = para.text.strip()
    if t: struct_text.append(t)
for table in doc1.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        struct_text.append(' | '.join(cells))
output['structured_dataset'] = '\n'.join(struct_text)
print(f"段落+表格行: {len(struct_text)}")

# 3. 解析游览指南docx
print("\n=== 解析游览指南docx ===")
doc2 = Document(r'd:\ai_guide\灵山胜境：历史、文化、景点特色与个性化游览指南.docx')
guide_text = []
for para in doc2.paragraphs:
    t = para.text.strip()
    if t: guide_text.append(t)
for table in doc2.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        guide_text.append(' | '.join(cells))
output['tour_guide'] = '\n'.join(guide_text)
print(f"段落+表格行: {len(guide_text)}")

# 保存为JSON
with open(r'd:\ai_guide\scripts\lingshan_data.json', 'w', encoding='utf-8') as f:
    json.dump(output, f, ensure_ascii=False, indent=2)
print(f"\n数据已保存: scripts/lingshan_data.json ({len(json.dumps(output, ensure_ascii=False))} 字符)")
